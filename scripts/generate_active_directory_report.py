from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("ActiveDirectory_Project_Report.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
    else:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)

    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9.5)

    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_code(doc, text):
    for line in text.strip("\n").split("\n"):
        para = doc.add_paragraph(style="Code")
        run = para.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)

    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10.5)

    body_para = cell.add_paragraph(body)
    body_para.paragraph_format.space_after = Pt(0)
    for run in body_para.runs:
        run.font.size = Pt(10)

    set_table_width(table, [9360])
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def paragraph(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.10
    return para


def section(doc, title):
    doc.add_heading(title, level=1)


def subsection(doc, title):
    doc.add_heading(title, level=2)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Code" not in doc.styles:
        style = doc.styles.add_style("Code", 1)
    else:
        style = doc.styles["Code"]
    style.font.name = "Consolas"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    style.font.size = Pt(9)
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def add_header_footer(doc):
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.text = "Engineering of Windows Server OS 2026 - Active Directory Project"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Active Directory Project Report | mydomain.com")
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Engineering of Windows Server OS 2026")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Active Directory Project Report")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = DARK_BLUE

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("A practical report on building and administering a Windows Server domain")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    add_table(
        doc,
        ["Item", "Value"],
        [
            ("Lab domain", "mydomain.com"),
            ("Domain Controller", "Windows Server 2019 or newer, named DC"),
            ("Client computer", "Windows client VM named CLIENT1"),
            ("Internal network", "VMnet0 / 172.16.0.0/24"),
            ("Core services", "AD DS, DNS, DHCP, RAS/NAT, Group Policy"),
        ],
        [2400, 6960],
    )
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    items = [
        "1. Introduction to Windows Server",
        "2. Introduction to Active Directory",
        "3. Lab Environment and Domain Setup",
        "4. User Account Configuration",
        "5. Groups Configuration",
        "6. File and Folder Permissions and Network Shares",
        "7. Group Policies for Remote Configuration",
        "8. Group Policies for Software Installation",
        "9. DNS Service Configuration",
        "10. Conclusion and Learning Outcomes",
        "Appendix A. Useful Administrative Commands",
    ]
    for item in items:
        para = doc.add_paragraph(item)
        para.paragraph_format.space_after = Pt(3)
    doc.add_page_break()


def build_report():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_toc(doc)

    section(doc, "1. Introduction to Windows Server")
    paragraph(doc, "Windows Server is Microsoft's server operating system platform for running business network services. It is different from desktop Windows because it is designed to serve many users and computers at the same time. Instead of being mainly a personal workstation, it becomes a central place where an organization can host identity services, file shares, DNS, DHCP, remote access, web applications, print services, storage, and management tools.")
    paragraph(doc, "In a company, the server is often the stable point of the network. Users may change desks, laptops may be replaced, and departments may grow, but the server keeps the shared rules and shared resources consistent. Windows Server provides this consistency through roles and features. A role is a major service such as Active Directory Domain Services, DNS Server, DHCP Server, File and Storage Services, or Remote Access. Features add supporting capabilities such as management tools, PowerShell modules, or security components.")
    paragraph(doc, "For this project, Windows Server is the foundation on which the Active Directory lab is built. The server VM becomes the Domain Controller, which means it stores the domain database, authenticates users, applies policies, provides DNS for the domain, leases client IP addresses through DHCP, and routes the internal client network to the internet through NAT.")
    add_table(doc, ["Server role", "Purpose in this project"], [
        ("Active Directory Domain Services", "Creates the domain, stores users, groups, computers, and security information."),
        ("DNS Server", "Lets clients locate the domain and the Domain Controller by name."),
        ("DHCP Server", "Automatically assigns IP settings to client machines on the internal network."),
        ("Routing and Remote Access / NAT", "Allows internal VMs to reach the internet through the Domain Controller."),
        ("File and Storage Services", "Hosts shared folders and demonstrates access control."),
        ("Group Policy Management", "Centrally configures users and computers across the domain."),
    ], [2600, 6760])
    add_callout(doc, "Key idea", "Windows Server is not just an operating system; in a domain environment it becomes the control plane for identity, access, networking, and configuration.")

    section(doc, "2. Introduction to Active Directory")
    paragraph(doc, "Active Directory Domain Services (AD DS) is Microsoft's directory service for Windows domain networks. A directory service is a structured database of network objects: users, groups, computers, printers, organizational units, and policies. The important point is that AD DS does not only store names; it also stores security information and relationships. This is why Active Directory can answer questions such as: Who is this user? Which groups do they belong to? Which computers are part of the domain? Which resources is this user allowed to access?")
    paragraph(doc, "In a workgroup, every computer manages its own local users. That is acceptable for two or three computers, but it becomes painful in a real organization. If a user changes password, leaves the company, or needs access to a folder, an administrator would need to touch many separate machines. A domain solves this by centralizing identity. The user logs in once with a domain account, and the Domain Controller authenticates the user for the whole network.")
    paragraph(doc, "A Domain Controller is a Windows Server with AD DS installed and promoted. Promotion means the server is no longer only a member server; it now hosts the directory database and participates in domain authentication. In this lab, the domain is treated as mydomain.com. When CLIENT1 joins this domain, it trusts the Domain Controller for logon, policy, DNS, and resource authorization.")
    subsection(doc, "Core terminology")
    add_table(doc, ["Term", "Meaning"], [
        ("Domain", "A security and administrative boundary such as mydomain.com."),
        ("Forest", "The top-level Active Directory structure. It can contain one or more domains."),
        ("Domain Controller", "A server that stores AD DS data and authenticates domain users."),
        ("Organizational Unit (OU)", "A container used to organize users and computers and apply Group Policy."),
        ("Group Policy Object (GPO)", "A set of user or computer configuration rules linked to sites, domains, or OUs."),
        ("LDAP", "Protocol used to query and modify directory objects."),
        ("Kerberos", "Default AD authentication protocol using tickets rather than repeatedly sending passwords."),
    ], [2400, 6960])
    paragraph(doc, "Active Directory is powerful because identity, authorization, and configuration all meet in one system. A user account can be placed into a group; that group can be granted access to a shared folder; a GPO can be linked to the user's OU; and the DNS service helps the computer find all of these domain services automatically.")

    section(doc, "3. Lab Environment and Domain Setup")
    paragraph(doc, "The lab uses virtualization so the whole domain can run on one physical computer. The Domain Controller VM needs two network adapters. The first adapter is NAT and connects outward through the host computer's home network, so the server can reach the internet. The second adapter is an internal network adapter on VMnet0, so the Domain Controller can communicate privately with client VMs. This separation is important: the Windows client should not receive an address directly from the home router; it should receive its internal network settings from the Domain Controller.")
    add_table(doc, ["Component", "Configuration"], [
        ("DC external adapter", "NAT adapter, identified as Ethernet0 in the lab instructions."),
        ("DC internal adapter", "VMnet0 internal network, static IP 172.16.0.1 / 255.255.255.0."),
        ("Internal gateway", "172.16.0.1, because the Domain Controller routes internal traffic."),
        ("DC preferred DNS", "127.0.0.1 loopback after DNS is installed, because the DC hosts the domain DNS zone."),
        ("Server name", "Rename the Windows Server VM to DC before promoting it."),
        ("Client adapter", "Internal VMnet0 only, so it depends on DHCP/RAS/NAT from the DC."),
    ], [2600, 6760])
    subsection(doc, "Practical setup sequence")
    add_numbered(doc, [
        "Install Windows Server 2019 or newer in the server VM.",
        "Configure two network adapters: NAT for internet access and VMnet0 for the internal domain network.",
        "Rename the adapters to clear names such as INTERNET and INTERNAL so later DHCP/RAS settings are not confused.",
        "Assign 172.16.0.1/24 to the internal adapter and do not configure a default gateway on that internal interface.",
        "Rename the computer to DC and restart.",
        "Install Active Directory Domain Services and promote the server to a new domain named mydomain.com.",
        "After restart, create a dedicated domain administrator account and add it to Domain Admins.",
        "Install and configure Routing and Remote Access with NAT so internal clients can reach the internet.",
        "Install DHCP Server and create a scope, for example 172.16.0.100-172.16.0.200, with router/DNS set to 172.16.0.1.",
        "Create CLIENT1 as a separate Windows client VM, attach it to VMnet0, obtain an IP lease, and join it to mydomain.com.",
    ])
    add_callout(doc, "Verification point", "On CLIENT1, ipconfig should show an address from the DC DHCP scope, gateway 172.16.0.1, and DNS 172.16.0.1. A successful ping to mydomain.com confirms that client DNS can resolve the domain name to the Domain Controller.")

    section(doc, "4. User Account Configuration")
    paragraph(doc, "User accounts are the most visible part of Active Directory. Each account represents a person or service that needs to authenticate. A domain user account can log on to domain-joined computers and can be granted access to network resources. In Active Directory Users and Computers, administrators can create accounts manually, set usernames, assign temporary passwords, require password changes at next logon, disable accounts, and move users into the correct OU.")
    paragraph(doc, "For a small number of users, the GUI is understandable and safe. The administrator opens Active Directory Users and Computers, navigates to an OU such as Users or Students, right-clicks, chooses New User, fills in the name and logon name, sets an initial password, and chooses account options. This method is useful when learning because it shows the object structure directly.")
    paragraph(doc, "For many users, PowerShell is better. It reduces repetitive work and avoids mistakes caused by clicking through the same wizard many times. With the ActiveDirectory module, the New-ADUser command can create accounts from a CSV file, place them in an OU, set initial passwords, and enable the accounts automatically.")
    add_code(doc, '''
Import-Module ActiveDirectory

New-ADUser -Name "Ali Demir" `
  -GivenName "Ali" `
  -Surname "Demir" `
  -SamAccountName "ademir" `
  -UserPrincipalName "ademir@mydomain.com" `
  -Path "OU=Students,DC=mydomain,DC=com" `
  -AccountPassword (ConvertTo-SecureString "P@ssw0rd2026!" -AsPlainText -Force) `
  -Enabled $true `
  -ChangePasswordAtLogon $true
''')
    paragraph(doc, "A good account design also includes lifecycle management. New users should be created in the correct OU, added to the correct groups, and forced to change temporary passwords. Users who leave the organization should usually be disabled first, not deleted immediately, because their account may still be needed for audit trails, mailbox recovery, or ownership review.")
    add_bullets(doc, [
        "Use clear naming standards, such as first initial plus surname or student ID format.",
        "Avoid giving administrative privileges to normal daily-use accounts.",
        "Use password and lockout policies through Group Policy rather than configuring each account separately.",
        "Disable accounts that are no longer active and document why they were disabled.",
    ])

    section(doc, "5. Groups Configuration")
    paragraph(doc, "Groups are how Active Directory scales access control. Instead of assigning permissions to each user one by one, administrators put users into groups and assign permissions to the group. This makes the system easier to understand and easier to audit. If a new student joins the class, they can be added to the Students group and immediately receive the same access as everyone else in that group.")
    paragraph(doc, "There are two main group categories: security groups and distribution groups. Security groups are used for permissions and can also be used for email distribution. Distribution groups are mainly for messaging and are not used to assign file permissions. For this project, security groups are the important type because they control access to folders, shares, and sometimes policy targeting.")
    add_table(doc, ["Group scope", "Best use"], [
        ("Global", "Collect users with similar roles from the same domain, such as GG_Students or GG_IT_Admins."),
        ("Domain Local", "Assign permissions to resources in the domain, such as DL_ProjectShare_Read."),
        ("Universal", "Used across multiple domains in larger forests; less important in this single-domain lab."),
    ], [2200, 7160])
    paragraph(doc, "A common best-practice pattern is AGDLP: Accounts go into Global groups; Global groups go into Domain Local groups; Domain Local groups receive Permissions. For example, user accounts are placed in GG_Students, GG_Students is placed into DL_CourseShare_Modify, and DL_CourseShare_Modify is granted Modify permission on a shared folder. This keeps people, roles, and resource permissions separate.")
    add_code(doc, '''
New-ADGroup -Name "GG_Students" -GroupScope Global -GroupCategory Security `
  -Path "OU=Groups,DC=mydomain,DC=com"

New-ADGroup -Name "DL_CourseShare_Modify" -GroupScope DomainLocal -GroupCategory Security `
  -Path "OU=Groups,DC=mydomain,DC=com"

Add-ADGroupMember -Identity "GG_Students" -Members "ademir"
Add-ADGroupMember -Identity "DL_CourseShare_Modify" -Members "GG_Students"
''')

    section(doc, "6. File and Folder Permissions and Network Shares")
    paragraph(doc, "File sharing is one of the most practical reasons to build a domain. A server can host a folder once and make it available to authorized users across the network. In Windows Server, two permission layers matter: share permissions and NTFS permissions. Share permissions apply when users access the folder over the network path, such as \\\\DC\\CourseShare. NTFS permissions apply directly to the folder and files on disk, whether access is local or remote.")
    paragraph(doc, "When both layers apply, Windows uses the most restrictive effective permission. For example, if the share permission allows Full Control but NTFS only allows Read, the user effectively gets Read. For this reason, many administrators keep share permissions broad, such as Authenticated Users or a resource group with Change permission, and then use NTFS permissions for precise control. The important rule is to be intentional and test the final effective access.")
    add_table(doc, ["Permission", "Meaning"], [
        ("Read", "Open and copy files, but not modify or delete them."),
        ("Modify", "Read, create, edit, and delete files and folders."),
        ("Full Control", "Modify permissions and take ownership in addition to all Modify rights."),
        ("Deny", "Explicitly blocks access, usually avoided unless there is a specific reason."),
    ], [2200, 7160])
    subsection(doc, "Example implementation")
    add_numbered(doc, [
        "Create a folder on the server, for example C:\\Shares\\CourseShare.",
        "Share it as CourseShare and choose controlled share permissions.",
        "Create groups such as DL_CourseShare_Read and DL_CourseShare_Modify.",
        "Assign NTFS Read permission to the read group and Modify permission to the modify group.",
        "Add role groups or user groups into the correct permission group.",
        "Test from CLIENT1 using a domain user account and the UNC path \\\\DC\\CourseShare.",
    ])
    paragraph(doc, "This is also where the difference between authentication and authorization becomes clear. Authentication proves the user is who they claim to be. Authorization decides what that authenticated user may do. Active Directory handles identity; group membership and NTFS/share permissions decide access.")

    section(doc, "7. Group Policies for Remote Configuration")
    paragraph(doc, "Group Policy is the centralized configuration system of Active Directory. A Group Policy Object contains settings that can apply to users or computers. When a GPO is linked to an OU, domain, or site, the objects in that container receive the settings. This means administrators can configure many machines without physically visiting each one.")
    paragraph(doc, "Group Policy has two major sides: Computer Configuration and User Configuration. Computer Configuration applies to the machine regardless of who logs in. It is useful for firewall rules, Windows Update behavior, security settings, scripts, and software settings. User Configuration follows the user and is useful for desktop restrictions, mapped drives, folder redirection, Control Panel restrictions, and logon scripts.")
    add_table(doc, ["GPO example", "Configuration target"], [
        ("Password policy", "Domain-level security settings for password length, complexity, and lockout."),
        ("Mapped network drive", "User Configuration preference mapping a drive letter to \\\\DC\\CourseShare."),
        ("Control Panel restriction", "User setting that hides or blocks specific Control Panel items."),
        ("Firewall rule", "Computer setting applied consistently to domain computers."),
        ("Wallpaper or desktop setting", "User setting that standardizes the environment for a department or lab."),
    ], [2600, 6760])
    paragraph(doc, "GPO processing order is often summarized as LSDOU: Local, Site, Domain, OU. Later policies can override earlier policies if settings conflict. In a simple lab, the most common links are at the domain level or an OU level. OU-level linking is usually cleaner because it targets only the intended users or computers.")
    add_code(doc, '''
gpupdate /force
gpresult /r
gpresult /h C:\\Temp\\gpo-report.html
''')
    paragraph(doc, "The gpupdate command asks the client to refresh policy immediately, while gpresult shows which GPOs actually applied. These commands are important because a GPO existing in the console does not guarantee that it reached a client. Scope, security filtering, WMI filtering, OU placement, DNS, and replication can all affect the result.")

    section(doc, "8. Group Policies for Software Installation")
    paragraph(doc, "Group Policy can also deploy software, especially MSI packages. The basic idea is that the administrator places the installer in a shared folder that domain computers can read, creates or edits a GPO, and configures Software Installation under Computer Configuration or User Configuration. When the policy applies, the client installs the software automatically according to the deployment mode.")
    paragraph(doc, "Software deployment through GPO works best with MSI installers because Windows Installer understands installation state, repair, upgrade, and removal. EXE installers are less straightforward and often require scripts or a more advanced deployment tool. In a small lab, MSI deployment is enough to demonstrate the principle of centralized software management.")
    add_table(doc, ["Deployment mode", "Description"], [
        ("Assigned to computers", "The software installs for the machine, commonly during startup before user logon."),
        ("Assigned to users", "The software is advertised to the user and installs when used or at logon depending on settings."),
        ("Published to users", "The software appears as available for installation, but is not forced automatically."),
    ], [2600, 6760])
    subsection(doc, "Safe deployment workflow")
    add_numbered(doc, [
        "Create a software share such as \\\\DC\\Software and grant read access to Domain Computers or Authenticated Users.",
        "Place the MSI installer in the shared folder and always reference it through the UNC path, not a local path like C:\\.",
        "Create a test OU containing CLIENT1 before linking the GPO broadly.",
        "Configure Software Installation in the GPO and assign the MSI package.",
        "Restart the client or run gpupdate, then verify installation and Event Viewer logs.",
    ])
    paragraph(doc, "Testing matters. A broken software GPO can slow down many machines or repeatedly fail at startup. In professional environments, modern tools such as Microsoft Intune, Configuration Manager, or third-party endpoint management platforms often replace or supplement GPO software deployment, but GPO deployment remains a useful foundation for understanding centralized installation.")

    section(doc, "9. DNS Service Configuration")
    paragraph(doc, "DNS is essential for Active Directory. Many people think DNS only resolves internet names like microsoft.com, but in a Windows domain DNS also helps clients find domain services. When a computer joins a domain or a user logs in, the client must locate a Domain Controller. It does this by querying DNS for special service records, especially SRV records for LDAP and Kerberos.")
    paragraph(doc, "In this lab, DNS is installed with AD DS and hosts the mydomain.com zone. The Domain Controller registers records for itself, including host records and service records. CLIENT1 receives 172.16.0.1 as its DNS server through DHCP. This is correct because CLIENT1 must ask the Domain Controller for internal domain names. If CLIENT1 used a public DNS server such as 8.8.8.8 as its only DNS server, it would not know where mydomain.com or the Domain Controller services are.")
    add_table(doc, ["Record type", "Purpose"], [
        ("A record", "Maps a hostname to an IPv4 address, for example DC.mydomain.com to 172.16.0.1."),
        ("AAAA record", "Maps a hostname to an IPv6 address."),
        ("CNAME record", "Creates an alias name pointing to another hostname."),
        ("SRV record", "Advertises service locations such as LDAP and Kerberos for Active Directory."),
        ("PTR record", "Reverse lookup record mapping an IP address back to a hostname."),
    ], [2200, 7160])
    subsection(doc, "Why the Domain Controller uses loopback DNS")
    paragraph(doc, "The Domain Controller can use 127.0.0.1 as its preferred DNS server because 127.0.0.1 means this same machine. Since the DC hosts the DNS zone for mydomain.com, asking itself is correct. External DNS servers do not know about the private lab domain, so the DC must use its own DNS service for domain records. For internet names, DNS forwarding can send unresolved external queries to an upstream resolver.")
    add_code(doc, '''
nslookup mydomain.com
nslookup -type=SRV _ldap._tcp.dc._msdcs.mydomain.com
ipconfig /all
ping mydomain.com
''')
    paragraph(doc, "DNS troubleshooting is one of the most important Active Directory skills. If DNS is wrong, domain join can fail, logon can become slow, Group Policy may not apply, and clients may be unable to find the Domain Controller. A healthy AD lab usually shows CLIENT1 using the DC as DNS, resolving mydomain.com, and receiving replies from 172.16.0.1.")

    section(doc, "10. Conclusion and Learning Outcomes")
    paragraph(doc, "This project demonstrates how Windows Server and Active Directory work together to create a centrally managed domain environment. The Domain Controller is not only a login server; it also becomes the DNS authority for the domain, the DHCP provider for clients, the policy source for configuration, and the routing point between the internal lab and the outside network.")
    paragraph(doc, "The practical lab begins with network design: one NAT adapter for the server's external connectivity and one internal adapter for domain communication. After AD DS promotion, the server becomes the authority for mydomain.com. DHCP gives CLIENT1 the correct IP configuration, DNS lets CLIENT1 find the domain, and domain join proves that authentication and name resolution are working.")
    paragraph(doc, "The administration topics build on that foundation. Users represent identities; groups make access scalable; folder permissions enforce authorization; Group Policy centralizes settings; software deployment shows how configuration can become automated; and DNS makes the entire domain discoverable. Together, these topics form the core skill set for Windows system administration.")
    add_callout(doc, "Main takeaway", "Active Directory is valuable because it turns many separate computers into one managed environment. Instead of configuring every machine and every user manually, the administrator manages identities, groups, permissions, policies, and services from the domain.")

    section(doc, "Appendix A. Useful Administrative Commands")
    add_table(doc, ["Command", "Use"], [
        ("ipconfig /all", "Check client IP address, DNS server, gateway, and domain suffix."),
        ("ping mydomain.com", "Verify that the domain name resolves and the DC responds."),
        ("nslookup mydomain.com", "Test DNS name resolution directly."),
        ("gpupdate /force", "Force a Group Policy refresh on a client."),
        ("gpresult /r", "Show which user and computer GPOs applied."),
        ("Get-ADUser -Filter *", "List AD users through PowerShell."),
        ("Get-ADGroupMember GG_Students", "Check group membership."),
        ("Restart-Computer", "Restart after configuration changes when needed."),
    ], [2800, 6560])
    paragraph(doc, "These commands are not a replacement for understanding the graphical consoles, but they make verification faster. In real administration, the best workflow is to configure carefully, test from a client computer, and document the result.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
